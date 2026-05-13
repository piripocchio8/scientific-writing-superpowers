"""sws_read_docx.py — python-docx wrapper used by drafting agents.

Cycle #7 hotfix (D22): native Read tool fails on .docx, so SWS ships its own
wrapper. Tests build a small fixture with python-docx itself in the test
setup; no static fixture file in the repo.
"""
from __future__ import annotations
import subprocess
import sys
from pathlib import Path
import pytest

REPO = Path(__file__).resolve().parents[1]
SCRIPT = REPO / "scripts" / "sws_read_docx.py"


def _make_fixture(tmp_path: Path) -> Path:
    """Build a small .docx with headings + body paragraphs + a styled run."""
    from docx import Document

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Peptide therapeutics have undergone a revival.")
    doc.add_paragraph("Semaglutide proved oral peptides are commercially viable.")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("Activity at 1 microM was reproducible.")
    doc.add_paragraph("All compounds showed >70% cell viability.")
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph("These data support the mechanism proposed earlier.")

    fixture = tmp_path / "fixture.docx"
    doc.save(str(fixture))
    return fixture


def _run(args, expect_zero: bool = True):
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True, text=True
    )
    if expect_zero:
        assert cp.returncode == 0, cp.stderr
    return cp


def test_full_read_prints_every_paragraph(tmp_path):
    fixture = _make_fixture(tmp_path)
    cp = _run([str(fixture)])
    out = cp.stdout
    assert "Introduction" in out
    assert "Peptide therapeutics have undergone a revival." in out
    assert "Semaglutide proved oral peptides are commercially viable." in out
    assert "Results" in out
    assert "Activity at 1 microM was reproducible." in out
    assert "Discussion" in out
    assert "These data support the mechanism proposed earlier." in out


def test_section_scoped_read_returns_only_target_section(tmp_path):
    fixture = _make_fixture(tmp_path)
    cp = _run([str(fixture), "--section", "Results"])
    out = cp.stdout
    assert "Activity at 1 microM was reproducible." in out
    assert "All compounds showed >70% cell viability." in out
    # Body paragraphs from other sections must not leak in.
    assert "Peptide therapeutics have undergone a revival." not in out
    assert "These data support the mechanism proposed earlier." not in out


def test_section_match_is_case_insensitive_substring(tmp_path):
    fixture = _make_fixture(tmp_path)
    cp = _run([str(fixture), "--section", "intro"])
    out = cp.stdout
    assert "Peptide therapeutics have undergone a revival." in out
    assert "Activity at 1 microM was reproducible." not in out


def test_paragraph_range_returns_only_that_slice(tmp_path):
    fixture = _make_fixture(tmp_path)
    # Paragraphs 1-indexed: 1=Intro heading, 2/3=intro body,
    # 4=Results heading, 5/6=results body, 7=Discussion heading, 8=disc body.
    cp = _run([str(fixture), "--paragraphs", "5-6"])
    out = cp.stdout
    assert "Activity at 1 microM was reproducible." in out
    assert "All compounds showed >70% cell viability." in out
    assert "Peptide therapeutics" not in out
    assert "These data support" not in out


def test_with_styles_prefixes_each_line(tmp_path):
    fixture = _make_fixture(tmp_path)
    cp = _run([str(fixture), "--with-styles"])
    out = cp.stdout
    # Heading paragraphs carry a Heading style.
    assert "[Heading 1] Introduction" in out
    # Body paragraphs use the Normal style by default.
    assert "[Normal] Peptide therapeutics have undergone a revival." in out


def test_missing_file_exits_2(tmp_path):
    cp = _run([str(tmp_path / "does_not_exist.docx")], expect_zero=False)
    assert cp.returncode == 2
    assert "not found" in cp.stderr.lower() or "no such" in cp.stderr.lower()


def test_malformed_file_exits_3(tmp_path):
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"this is not a real docx file")
    cp = _run([str(bad)], expect_zero=False)
    assert cp.returncode == 3
