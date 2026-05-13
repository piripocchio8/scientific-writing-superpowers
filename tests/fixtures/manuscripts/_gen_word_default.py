"""Generator for tests/fixtures/manuscripts/word_default.docx.

Run once to regenerate the committed binary:
  python3 tests/fixtures/manuscripts/_gen_word_default.py

Produces a .docx with Word built-in styles (Heading 1, Heading 2, Normal) and
one paragraph containing an explicit italic run, so tests/test_restyle_docx.py
can verify that restyle preserves direct character formatting (R2 mitigation).
"""
from __future__ import annotations

from pathlib import Path


def generate() -> None:
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()

    # Heading 1 → should map to SWS-H1
    doc.add_heading("Introduction", level=1)

    # Heading 2 → should map to SWS-H2
    doc.add_heading("Background", level=2)

    # Normal paragraph with one italic run → italic must survive restyle
    para = doc.add_paragraph()
    para.add_run("Ordinary text followed by ")
    italic_run = para.add_run("italic text")
    italic_run.italic = True
    para.add_run(" and back to normal.")

    # Normal paragraph to test plain body mapping
    doc.add_paragraph(
        "A second body paragraph without any direct character formatting."
    )

    # Paragraph with highlight colour (yellow) → highlight must survive restyle
    highlight_para = doc.add_paragraph()
    highlighted_run = highlight_para.add_run("Highlighted passage.")
    # Use XML-level highlight (WD_COLOR_INDEX constants are in python-docx)
    from docx.oxml.ns import qn
    from lxml import etree
    rPr = highlighted_run._r.get_or_add_rPr()
    highlight_el = etree.SubElement(rPr, qn("w:highlight"))
    highlight_el.set(qn("w:val"), "yellow")

    # A References heading followed by a reference entry
    doc.add_heading("References", level=1)
    doc.add_paragraph(
        "[1] Smith, J.; Doe, A. J. Am. Chem. Soc. 2023, 145, 1234–1240."
    )

    out = Path(__file__).parent / "word_default.docx"
    doc.save(str(out))
    print(f"Generated {out}")


if __name__ == "__main__":
    generate()
