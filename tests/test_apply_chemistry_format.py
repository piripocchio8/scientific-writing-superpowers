"""sws_apply_chemistry_format.py — pytest cases per plan Task 1D Step 2.

Covers:
- "H2O is a solvent" → subscript on the 2.
- "et al." → italic.
- "E. coli grows" → italic on "E. coli".
- Genus species pattern → severity=suggest, NOT applied with default severity.
- "Figure 1. Caption" → bold on "Figure 1.".
- format=latex marker → no-op exit 0.
- --dry-run does not write the output file.
- Idempotence: apply twice yields same run count.
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[1]
SCRIPT = REPO / "scripts" / "sws_apply_chemistry_format.py"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _run(args: list[str], env=None, expect_zero: bool = True) -> subprocess.CompletedProcess:
    import os
    run_env = os.environ.copy()
    if env:
        run_env.update(env)
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True,
        text=True,
        env=run_env,
    )
    if expect_zero:
        assert cp.returncode == 0, f"stdout={cp.stdout!r}\nstderr={cp.stderr!r}"
    return cp


def _make_docx(tmp_path: Path, paragraphs: list[str]) -> Path:
    """Build a minimal .docx with the given paragraph texts."""
    from docx import Document
    doc = Document()
    # Remove the default blank paragraph.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for text in paragraphs:
        doc.add_paragraph(text)
    out = tmp_path / "input.docx"
    doc.save(str(out))
    return out


def _get_runs(docx_path: Path) -> list[dict]:
    """Return list of {text, italic, bold, subscript, superscript} for every run."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    results = []
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn("w:rPr"))
            subscript = False
            superscript = False
            if rPr is not None:
                va = rPr.find(qn("w:vertAlign"))
                if va is not None:
                    val = va.get(qn("w:val"))
                    subscript = val == "subscript"
                    superscript = val == "superscript"
            results.append({
                "text": run.text,
                "italic": bool(run.italic),
                "bold": bool(run.bold),
                "subscript": subscript,
                "superscript": superscript,
            })
    return results


def _count_runs_with(docx_path: Path, **kwargs) -> int:
    """Count runs matching ALL the given kwargs."""
    runs = _get_runs(docx_path)
    count = 0
    for r in runs:
        if all(r.get(k) == v for k, v in kwargs.items()):
            count += 1
    return count


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_h2o_subscript(tmp_path):
    """'H2O is a solvent' → the '2' run gets subscript formatting."""
    src = _make_docx(tmp_path, ["H2O is a solvent"])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out)])
    subscript_runs = _count_runs_with(out, subscript=True)
    assert subscript_runs >= 1, "Expected at least one subscript run for '2' in H2O"
    # Confirm the subscripted text is '2'.
    runs = _get_runs(out)
    subscripted_texts = [r["text"] for r in runs if r["subscript"]]
    assert "2" in subscripted_texts, f"Subscript run should be '2', got: {subscripted_texts}"


def test_et_al_italic(tmp_path):
    """'Smith et al. reported' → 'et al.' run is italic."""
    src = _make_docx(tmp_path, ["Smith et al. reported"])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out)])
    runs = _get_runs(out)
    italic_texts = [r["text"] for r in runs if r["italic"]]
    assert any("et al" in t for t in italic_texts), (
        f"Expected 'et al.' to be italic; italic runs: {italic_texts}"
    )


def test_e_coli_italic(tmp_path):
    """'E. coli grows' → 'E. coli' span is italic (species_abbreviated, auto)."""
    src = _make_docx(tmp_path, ["E. coli grows"])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out)])
    runs = _get_runs(out)
    italic_texts = [r["text"] for r in runs if r["italic"]]
    assert italic_texts, f"Expected italic runs; all runs: {runs}"
    combined_italic = "".join(italic_texts)
    assert "coli" in combined_italic, (
        f"Expected 'coli' in italic span; italic texts: {italic_texts}"
    )


def test_genus_species_suggest_not_applied_by_default(tmp_path):
    """'Staphylococcus aureus biofilm' — genus_species is suggest; NOT applied by default."""
    src = _make_docx(tmp_path, ["Staphylococcus aureus biofilm"])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out)])
    # With default severity=auto, species_names (all suggest) should NOT be italicised.
    runs = _get_runs(out)
    italic_texts = [r["text"] for r in runs if r["italic"]]
    # The full text should remain unitalicised.
    combined_italic = "".join(italic_texts)
    assert "Staphylococcus" not in combined_italic and "aureus" not in combined_italic, (
        f"genus_species should not be applied with default severity; italic: {italic_texts}"
    )


def test_genus_species_suggest_applied_with_severity_all(tmp_path):
    """'Staphylococcus aureus biofilm' — genus_species applied when --severity all."""
    src = _make_docx(tmp_path, ["Staphylococcus aureus biofilm"])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out), "--severity", "all"])
    runs = _get_runs(out)
    italic_texts = [r["text"] for r in runs if r["italic"]]
    combined_italic = "".join(italic_texts)
    assert "aureus" in combined_italic or "Staphylococcus" in combined_italic, (
        f"Expected italic when --severity all; italic texts: {italic_texts}"
    )


def test_figure_caption_bold(tmp_path):
    """'Figure 1. Reaction scheme.' → 'Figure 1.' prefix is bold."""
    src = _make_docx(tmp_path, ["Figure 1. Reaction scheme."])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out)])
    runs = _get_runs(out)
    bold_texts = [r["text"] for r in runs if r["bold"]]
    assert bold_texts, f"Expected bold runs for figure label; runs: {runs}"
    combined_bold = "".join(bold_texts)
    assert "Figure" in combined_bold, (
        f"Expected 'Figure' in bold span; bold texts: {bold_texts}"
    )


def test_latex_marker_noop(tmp_path):
    """format=latex in project marker → skipped exit 0, output file NOT written."""
    src = _make_docx(tmp_path, ["H2O is a solvent"])
    out = tmp_path / "out.docx"
    # Write a marker with format: latex.
    marker = tmp_path / ".sws-project.local.md"
    marker.write_text("format: latex\n", encoding="utf-8")

    cp = _run(
        [str(src), "--out", str(out)],
        env={"SWS_MARKER": str(marker)},
    )
    assert "skipped" in cp.stdout.lower() or "latex" in cp.stdout.lower(), (
        f"Expected 'skipped (format=latex)' message; stdout={cp.stdout!r}"
    )
    # Output file should NOT have been written.
    assert not out.exists(), "Output file should not be written for format=latex"


def test_dry_run_no_output_file(tmp_path):
    """--dry-run does not write the output file."""
    src = _make_docx(tmp_path, ["H2O is a solvent"])
    out = tmp_path / "out.docx"
    _run([str(src), "--out", str(out), "--dry-run"])
    assert not out.exists(), "--dry-run should not write the output file"


def test_dry_run_stdout_reports_changes(tmp_path):
    """--dry-run prints a summary without writing."""
    src = _make_docx(tmp_path, ["H2O is a solvent", "Smith et al. reported"])
    out = tmp_path / "out.docx"
    cp = _run([str(src), "--out", str(out), "--dry-run"])
    # Should mention dry-run and counts.
    assert "dry-run" in cp.stdout.lower() or "modified" in cp.stdout.lower(), (
        f"Expected dry-run summary; stdout={cp.stdout!r}"
    )


def test_missing_input_exits_2(tmp_path):
    """Non-existent input → exit code 2."""
    cp = _run([str(tmp_path / "no_such_file.docx")], expect_zero=False)
    assert cp.returncode == 2


def test_idempotence(tmp_path):
    """Apply twice: run count with subscript/italic should not increase on second pass."""
    src = _make_docx(tmp_path, [
        "H2O is a solvent",
        "Smith et al. reported",
        "Figure 1. Reaction scheme.",
    ])
    pass1 = tmp_path / "pass1.docx"
    pass2 = tmp_path / "pass2.docx"
    _run([str(src), "--out", str(pass1)])
    _run([str(pass1), "--out", str(pass2)])

    runs1 = _get_runs(pass1)
    runs2 = _get_runs(pass2)

    sub1 = sum(1 for r in runs1 if r["subscript"])
    sub2 = sum(1 for r in runs2 if r["subscript"])
    assert sub2 == sub1, f"Subscript count changed on second pass: {sub1} → {sub2}"

    ital1 = sum(1 for r in runs1 if r["italic"])
    ital2 = sum(1 for r in runs2 if r["italic"])
    assert ital2 == ital1, f"Italic count changed on second pass: {ital1} → {ital2}"

    bold1 = sum(1 for r in runs1 if r["bold"])
    bold2 = sum(1 for r in runs2 if r["bold"])
    assert bold2 == bold1, f"Bold count changed on second pass: {bold1} → {bold2}"


def test_inplace_without_out(tmp_path):
    """Without --out, the script modifies the input file in place."""
    src = _make_docx(tmp_path, ["Smith et al. reported"])
    _run([str(src)])  # no --out
    # The file must still be a valid docx with italic runs.
    runs = _get_runs(src)
    italic_texts = [r["text"] for r in runs if r["italic"]]
    assert any("et al" in t for t in italic_texts), (
        f"In-place edit: expected italic 'et al.'; italic: {italic_texts}"
    )


def test_suggest_reported_in_stdout(tmp_path):
    """Suggest-severity matches are reported in stdout even at default severity."""
    src = _make_docx(tmp_path, ["Staphylococcus aureus biofilm"])
    out = tmp_path / "out.docx"
    cp = _run([str(src), "--out", str(out)])
    # The suggest pattern should appear in stdout as a suggestion.
    assert "suggest" in cp.stdout.lower(), (
        f"Expected suggest report in stdout; stdout={cp.stdout!r}"
    )
