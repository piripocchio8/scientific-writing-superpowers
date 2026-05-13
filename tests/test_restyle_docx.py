"""sws_restyle_docx.py — tests per plan Task 1C Step 2.

Covers:
- Word-default Heading 1/2 styles map to SWS-H1/SWS-H2.
- Direct italic run on a Normal paragraph survives restyle (R2 mitigation).
- Idempotence: restyle(restyle(x)) produces no style drift.
- Already-SWS-styled docx → no-op (styles unchanged).
- Highlight colour is preserved on the run after restyle.
- Missing input → exit code 2.
- python-docx parse error → exit code 3.
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[1]
SCRIPT = REPO / "scripts" / "sws_restyle_docx.py"
FIXTURE = REPO / "tests" / "fixtures" / "manuscripts" / "word_default.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _run(args: list[str], expect_zero: bool = True) -> subprocess.CompletedProcess:
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True,
        text=True,
    )
    if expect_zero:
        assert cp.returncode == 0, cp.stderr
    return cp


def _make_word_default(tmp_path: Path) -> Path:
    """Build a fresh Word-default-styled .docx for use in tests."""
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_heading("Background", level=2)

    # Normal paragraph with explicit italic run
    para = doc.add_paragraph()
    para.add_run("Plain text ")
    italic_run = para.add_run("italic word")
    italic_run.italic = True
    para.add_run(" end.")

    # Paragraph with highlight colour
    h_para = doc.add_paragraph()
    h_run = h_para.add_run("Highlighted.")
    rPr = h_run._r.get_or_add_rPr()
    hl = etree.SubElement(rPr, qn("w:highlight"))
    hl.set(qn("w:val"), "yellow")

    # References section
    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Author A. Journal 2023.")

    out = tmp_path / "word_default.docx"
    doc.save(str(out))
    return out


def _make_sws_styled(tmp_path: Path) -> Path:
    """Build a .docx that already uses SWS styles."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    for style_name, bold, italic, size in [
        ("SWS-H1", True, False, 12),
        ("SWS-H2", True, True, 12),
        ("SWS-Body", False, False, 12),
        ("SWS-Caption", False, False, 10),
        ("SWS-References", False, False, 10),
    ]:
        s = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.italic = italic

    p = doc.add_paragraph("A heading paragraph")
    p.style = doc.styles["SWS-H1"]

    p2 = doc.add_paragraph("Body text paragraph.")
    p2.style = doc.styles["SWS-Body"]

    out = tmp_path / "sws_styled.docx"
    doc.save(str(out))
    return out


def _paragraph_styles(path: Path) -> list[str]:
    """Return the list of paragraph style names from a .docx."""
    from docx import Document
    doc = Document(str(path))
    return [p.style.name for p in doc.paragraphs]


def _run_italic_survives(path: Path) -> bool:
    """Return True if any run in the document is marked italic at the run level."""
    from docx import Document
    doc = Document(str(path))
    for para in doc.paragraphs:
        for run in para.runs:
            if run.italic:
                return True
    return False


def _highlight_survives(path: Path) -> bool:
    """Return True if any run contains a w:highlight XML element."""
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(path))
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn("w:rPr"))
            if rPr is not None and rPr.find(qn("w:highlight")) is not None:
                return True
    return False


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_word_default_heading1_maps_to_sws_h1(tmp_path):
    """Heading 1 → SWS-H1 after restyle."""
    src = _make_word_default(tmp_path)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    styles = _paragraph_styles(out)
    assert "SWS-H1" in styles


def test_word_default_heading2_maps_to_sws_h2(tmp_path):
    """Heading 2 → SWS-H2 after restyle."""
    src = _make_word_default(tmp_path)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    styles = _paragraph_styles(out)
    assert "SWS-H2" in styles


def test_normal_paragraph_maps_to_sws_body(tmp_path):
    """Normal paragraphs → SWS-Body after restyle."""
    src = _make_word_default(tmp_path)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    styles = _paragraph_styles(out)
    assert "SWS-Body" in styles


def test_references_section_maps_to_sws_references(tmp_path):
    """Paragraphs under the References heading → SWS-References."""
    src = _make_word_default(tmp_path)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    styles = _paragraph_styles(out)
    assert "SWS-References" in styles


def test_italic_run_survives_restyle(tmp_path):
    """Explicit italic run on a Normal paragraph is preserved after restyle (R2)."""
    src = _make_word_default(tmp_path)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    assert _run_italic_survives(out), "Italic run was lost during restyle"


def test_highlight_colour_survives_restyle(tmp_path):
    """Highlight colour on a run is preserved after restyle."""
    src = _make_word_default(tmp_path)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    assert _highlight_survives(out), "Highlight colour was lost during restyle"


def test_idempotence(tmp_path):
    """restyle(restyle(x)) produces the same paragraph style list as restyle(x)."""
    src = _make_word_default(tmp_path)
    pass1 = tmp_path / "pass1.docx"
    pass2 = tmp_path / "pass2.docx"
    _run([str(src), "--out", str(pass1)])
    _run([str(pass1), "--out", str(pass2)])
    assert _paragraph_styles(pass1) == _paragraph_styles(pass2)


def test_already_sws_styled_is_noop(tmp_path):
    """A docx already using SWS styles comes out with the same style list."""
    src = _make_sws_styled(tmp_path)
    styles_before = _paragraph_styles(src)
    out = tmp_path / "restyled.docx"
    _run([str(src), "--out", str(out)])
    styles_after = _paragraph_styles(out)
    assert styles_before == styles_after


def test_missing_input_exits_2(tmp_path):
    """Non-existent input file → exit code 2."""
    cp = _run([str(tmp_path / "no_such_file.docx")], expect_zero=False)
    assert cp.returncode == 2


def test_malformed_file_exits_3(tmp_path):
    """Corrupt (non-docx) input → exit code 3."""
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"this is not a real docx")
    cp = _run([str(bad)], expect_zero=False)
    assert cp.returncode == 3


def test_committed_fixture_restyles_without_error():
    """The committed word_default.docx fixture restyles without error."""
    assert FIXTURE.is_file(), f"Fixture not found: {FIXTURE}"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), str(FIXTURE), "--out", "/dev/null"],
        capture_output=True,
        text=True,
    )
    assert cp.returncode == 0, cp.stderr


def test_caption_paragraph_maps_to_sws_caption(tmp_path):
    """Paragraphs starting with 'Figure N' or 'Table N' → SWS-Caption."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Figure 1. The structure of the compound.")
    doc.add_paragraph("Table 1. Summary of results.")
    doc.add_paragraph("Fig. 2. Another figure caption.")
    src = tmp_path / "captions.docx"
    doc.save(str(src))

    out = tmp_path / "captions_restyled.docx"
    _run([str(src), "--out", str(out)])
    styles = _paragraph_styles(out)
    assert styles.count("SWS-Caption") == 3
