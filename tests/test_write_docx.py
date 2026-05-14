"""sws_write_docx.py — tests for the SWS markdown→docx WRITE wrapper.

Cycle #8 Task 1B.  Mirrors the tmpdir/subprocess pattern of test_read_docx.py.
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[1]
SCRIPT = REPO / "scripts" / "sws_write_docx.py"
STYLE_REF = REPO / "references" / "docx-style.md"

ALL_SWS_STYLES = ["SWS-Body", "SWS-H1", "SWS-H2", "SWS-Caption", "SWS-References"]


def _run(args: list[str], expect_zero: bool = True) -> subprocess.CompletedProcess:
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True, text=True,
    )
    if expect_zero:
        assert cp.returncode == 0, f"stdout={cp.stdout!r}\nstderr={cp.stderr!r}"
    return cp


def _style_names(docx_path: Path) -> list[str]:
    """Return paragraph style names for every paragraph in the .docx."""
    from docx import Document
    doc = Document(str(docx_path))
    return [p.style.name for p in doc.paragraphs]


def _paragraph_texts(docx_path: Path) -> list[str]:
    from docx import Document
    doc = Document(str(docx_path))
    return [p.text for p in doc.paragraphs]


def _defined_style_names(docx_path: Path) -> list[str]:
    """Return names of all paragraph styles defined in the .docx."""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    doc = Document(str(docx_path))
    return [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]


# ---------------------------------------------------------------------------
# Empty markdown → empty docx with five styles defined
# ---------------------------------------------------------------------------

def test_empty_markdown_produces_empty_docx_with_sws_styles(tmp_path):
    md = tmp_path / "empty.md"
    md.write_text("", encoding="utf-8")
    out = tmp_path / "out.docx"
    _run([str(out), "--from-markdown", str(md)])
    assert out.is_file()
    defined = _defined_style_names(out)
    for sws in ALL_SWS_STYLES:
        assert sws in defined, f"{sws} not defined in docx styles"
    # No actual paragraphs.
    assert _paragraph_texts(out) == []


# ---------------------------------------------------------------------------
# H1 + body → SWS-H1 + SWS-Body
# ---------------------------------------------------------------------------

def test_h1_and_body_get_correct_styles(tmp_path):
    md = tmp_path / "h1body.md"
    md.write_text(
        "# Introduction\n"
        "Peptides have shown promise in drug delivery.\n"
        "Cell viability remained above 80% in all conditions.\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    _run([str(out), "--from-markdown", str(md)])
    styles = _style_names(out)
    assert styles[0] == "SWS-H1"
    assert all(s == "SWS-Body" for s in styles[1:])


# ---------------------------------------------------------------------------
# All five style types
# ---------------------------------------------------------------------------

_FIVE_STYLE_MD = """\
# Results

The yield was 87%.

## Sub-section

A deeper look at the data.

**Figure 1.** Reaction scheme showing the key intermediate.

## References

1. Smith J. et al. *Nature* 2024, 600, 123.
2. Jones A. et al. *JACS* 2023, 145, 456.
"""


def test_all_five_styles_applied_correctly(tmp_path):
    md = tmp_path / "all5.md"
    md.write_text(_FIVE_STYLE_MD, encoding="utf-8")
    out = tmp_path / "out.docx"
    _run([str(out), "--from-markdown", str(md)])

    styles = _style_names(out)
    texts = _paragraph_texts(out)

    pairs = list(zip(texts, styles))

    def find_style(text_fragment: str) -> str:
        for t, s in pairs:
            if text_fragment in t:
                return s
        raise AssertionError(f"text fragment not found: {text_fragment!r}")

    assert find_style("Results") == "SWS-H1"
    assert find_style("yield was") == "SWS-Body"
    assert find_style("Sub-section") == "SWS-H2"
    assert find_style("deeper look") == "SWS-Body"
    assert find_style("Figure 1.") == "SWS-Caption"
    # The References H2 heading is emitted as SWS-H1 (top-level section).
    assert find_style("References") == "SWS-H1"
    # Reference entries fall under SWS-References.
    assert find_style("Smith J.") == "SWS-References"
    assert find_style("Jones A.") == "SWS-References"


# ---------------------------------------------------------------------------
# Missing input → exit 2
# ---------------------------------------------------------------------------

def test_missing_input_file_exits_2(tmp_path):
    out = tmp_path / "out.docx"
    cp = _run([str(out), "--from-markdown", str(tmp_path / "no_such.md")], expect_zero=False)
    assert cp.returncode == 2
    assert "not found" in cp.stderr.lower()


# ---------------------------------------------------------------------------
# Style values sourced from references/docx-style.md
# ---------------------------------------------------------------------------

def test_style_values_read_from_style_ref(tmp_path, monkeypatch):
    """Replacing the YAML in docx-style.md changes the generated docx font."""
    # Write a patched style ref with font Courier New instead of Arial.
    patched_ref = tmp_path / "docx-style.md"
    original = STYLE_REF.read_text(encoding="utf-8")
    patched = original.replace("Arial", "Courier New")
    patched_ref.write_text(patched, encoding="utf-8")

    # Monkey-patch the module-level _STYLE_REF path by injecting an env override.
    # We call the script as a subprocess, so we do it via a wrapper that rewrites
    # the _STYLE_REF constant at import time using PYTHONPATH tricks.
    #
    # Simpler approach: import the module directly and patch _STYLE_REF.
    import importlib, types
    spec_obj = importlib.util.spec_from_file_location("sws_write_docx", str(SCRIPT))
    mod = importlib.util.module_from_spec(spec_obj)

    # Patch before exec_module so _STYLE_REF resolves to our patched file.
    monkeypatch.setattr("builtins.__import__", __import__)   # no-op guard

    # Override _STYLE_REF after loading.
    spec_obj.loader.exec_module(mod)
    mod._STYLE_REF = patched_ref

    style_defs = mod._load_style_defs()
    assert style_defs["SWS-Body"]["font"] == "Courier New"

    # Now build a docx and check the font on the SWS-Body style.
    md = tmp_path / "sample.md"
    md.write_text("# Title\nBody text here.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    lines = mod._read_markdown_file(md)
    specs = mod._classify_lines(lines)
    mod._build_docx(specs, style_defs, out)

    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    doc = Document(str(out))
    body_style = next(
        (s for s in doc.styles
         if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name == "SWS-Body"),
        None,
    )
    assert body_style is not None
    assert body_style.font.name == "Courier New"


# ---------------------------------------------------------------------------
# --from-stitched alias works identically to --from-markdown
# ---------------------------------------------------------------------------

def test_from_stitched_alias(tmp_path):
    md = tmp_path / "stitched.md"
    md.write_text("# Intro\nHello world.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    _run([str(out), "--from-stitched", str(md)])
    styles = _style_names(out)
    assert "SWS-H1" in styles
    assert "SWS-Body" in styles


# ---------------------------------------------------------------------------
# --from-drafts-dir requires --profile
# ---------------------------------------------------------------------------

def test_from_drafts_dir_without_profile_errors(tmp_path):
    drafts = tmp_path / "_drafts"
    drafts.mkdir()
    (drafts / "intro.md").write_text("# Intro\nBody.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    cp = _run([str(out), "--from-drafts-dir", str(drafts)], expect_zero=False)
    assert cp.returncode != 0


def test_from_drafts_dir_stitches_sections(tmp_path):
    drafts = tmp_path / "_drafts"
    drafts.mkdir()
    (drafts / "01_intro.md").write_text("# Introduction\nBody intro.\n", encoding="utf-8")
    (drafts / "02_results.md").write_text("# Results\nBody results.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    _run([str(out), "--from-drafts-dir", str(drafts), "--profile", "full-article"])
    styles = _style_names(out)
    assert styles.count("SWS-H1") == 2
    assert styles.count("SWS-Body") == 2
