"""Write a .docx from markdown — SWS WRITE wrapper around python-docx.

Cycle #8 D5 / D22 deferral from cycle #7.  Converts markdown source to a
.docx that uses the SWS custom style canon defined in references/docx-style.md.

Style values are read from the YAML frontmatter of references/docx-style.md at
start-up — do NOT hardcode font names, sizes, or weights here.

CLI:
  sws_write_docx.py <output.docx> --from-markdown <md-file>
  sws_write_docx.py <output.docx> --from-stitched <stitched.md>
  sws_write_docx.py <output.docx> --from-drafts-dir <drafts-dir> --profile <profile-id>

Exit codes:
  0  ok
  2  input file/directory not found
  3  python-docx or YAML error
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Any

_REPO = Path(__file__).resolve().parent.parent
_STYLE_REF = _REPO / "references" / "docx-style.md"

# Heading / caption detection regexes (against raw markdown line).
_RE_H1 = re.compile(r"^#\s+(.+)$")
_RE_H2 = re.compile(r"^##\s+(.+)$")
_RE_CAPTION = re.compile(r"^\*\*(Figure|Table)\s+\d+[.:)]?\*\*")

_HALF_PT = 2   # python-docx size unit is half-points; multiply by 2 per pt


# ---------------------------------------------------------------------------
# Style reference reader
# ---------------------------------------------------------------------------

def _load_style_defs() -> dict[str, dict[str, Any]]:
    """Parse the YAML frontmatter of references/docx-style.md.

    Returns the ``styles`` dict, e.g.:
        {'SWS-Body': {'font': 'Arial', 'size': 12, 'bold': False, ...}, ...}
    """
    try:
        import yaml
    except ImportError as exc:
        print(f"sws_write_docx: PyYAML is required: {exc}", file=sys.stderr)
        sys.exit(3)

    text = _STYLE_REF.read_text(encoding="utf-8")
    # Extract YAML between the first and second '---' fence.
    parts = text.split("---")
    if len(parts) < 3:
        print("sws_write_docx: docx-style.md has no YAML frontmatter", file=sys.stderr)
        sys.exit(3)
    try:
        meta = yaml.safe_load(parts[1])
    except yaml.YAMLError as exc:
        print(f"sws_write_docx: YAML parse error in docx-style.md: {exc}", file=sys.stderr)
        sys.exit(3)

    styles = meta.get("styles")
    if not styles:
        print("sws_write_docx: 'styles' key missing in docx-style.md", file=sys.stderr)
        sys.exit(3)
    return styles


# ---------------------------------------------------------------------------
# python-docx style injection
# ---------------------------------------------------------------------------

def _ensure_sws_styles(doc, style_defs: dict[str, dict[str, Any]]) -> None:
    """Add SWS paragraph styles to *doc* if they are not already present.

    python-docx's Document() starts with built-in Word styles.  We inject the
    five SWS custom styles (SWS-Body, SWS-H1, SWS-H2, SWS-Caption,
    SWS-References) sourced from style_defs.
    """
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

    for sws_name, props in style_defs.items():
        # Skip if already present (idempotent).
        if sws_name in [s.name for s in doc.styles]:
            continue

        style = doc.styles.add_style(sws_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = props.get("font", "Arial")
        size = props.get("size", 12)
        font.size = Pt(size)
        font.bold = bool(props.get("bold", False))
        font.italic = bool(props.get("italic", False))


# ---------------------------------------------------------------------------
# Markdown → paragraph classification
# ---------------------------------------------------------------------------

class _ParagraphSpec:
    """A single output paragraph: text + SWS style name."""
    __slots__ = ("text", "style")

    def __init__(self, text: str, style: str):
        self.text = text
        self.style = style


def _classify_lines(lines: list[str]) -> list[_ParagraphSpec]:
    """Map markdown lines to _ParagraphSpec entries.

    Rules (applied in order per non-blank line):
    1. ``## References`` → following lines are SWS-References.
    2. ``## <anything else>`` → SWS-H2, body resumes below.
    3. ``# <anything>`` → SWS-H1.
    4. Lines matching **Figure N.** / **Table N.** → SWS-Caption.
    5. Everything else → SWS-Body.

    Blank lines are skipped (Word paragraphs use spacing, not blank lines).
    """
    specs: list[_ParagraphSpec] = []
    in_references = False

    for raw in lines:
        line = raw.rstrip("\n")

        # Blank lines → skip.
        if not line.strip():
            continue

        # Check H2 first (two hashes), then H1 (one hash).
        m2 = _RE_H2.match(line)
        if m2:
            heading_text = m2.group(1).strip()
            in_references = heading_text.lower() == "references"
            if in_references:
                # Emit the References heading itself as SWS-H1 (top-level).
                specs.append(_ParagraphSpec(heading_text, "SWS-H1"))
            else:
                specs.append(_ParagraphSpec(heading_text, "SWS-H2"))
            continue

        m1 = _RE_H1.match(line)
        if m1:
            in_references = False
            specs.append(_ParagraphSpec(m1.group(1).strip(), "SWS-H1"))
            continue

        if in_references:
            specs.append(_ParagraphSpec(line, "SWS-References"))
            continue

        if _RE_CAPTION.match(line):
            # Strip bold markers for storage; python-docx styles own formatting.
            clean = line.replace("**", "")
            specs.append(_ParagraphSpec(clean, "SWS-Caption"))
            continue

        specs.append(_ParagraphSpec(line, "SWS-Body"))

    return specs


# ---------------------------------------------------------------------------
# Source readers
# ---------------------------------------------------------------------------

def _read_markdown_file(path: Path) -> list[str]:
    if not path.is_file():
        print(f"sws_write_docx: file not found: {path}", file=sys.stderr)
        sys.exit(2)
    return path.read_text(encoding="utf-8").splitlines(keepends=True)


def _read_drafts_dir(drafts_dir: Path, profile_id: str) -> list[str]:
    """Stitch _drafts/<section>.md files in sorted order for the given profile."""
    if not drafts_dir.is_dir():
        print(f"sws_write_docx: drafts directory not found: {drafts_dir}", file=sys.stderr)
        sys.exit(2)
    md_files = sorted(drafts_dir.glob("*.md"))
    if not md_files:
        print(f"sws_write_docx: no .md files in {drafts_dir}", file=sys.stderr)
        sys.exit(2)
    lines: list[str] = []
    for f in md_files:
        lines.extend(f.read_text(encoding="utf-8").splitlines(keepends=True))
        lines.append("\n")   # blank separator between sections
    return lines


# ---------------------------------------------------------------------------
# .docx builder
# ---------------------------------------------------------------------------

def _build_docx(specs: list[_ParagraphSpec], style_defs: dict[str, dict], output: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        print(f"sws_write_docx: python-docx is required: {exc}", file=sys.stderr)
        sys.exit(3)

    try:
        doc = Document()
        # Remove the blank default paragraph that python-docx adds.
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)

        _ensure_sws_styles(doc, style_defs)

        for spec in specs:
            p = doc.add_paragraph(spec.text, style=spec.style)

        doc.save(str(output))
    except Exception as exc:
        print(f"sws_write_docx: python-docx error: {exc}", file=sys.stderr)
        sys.exit(3)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        description="Write a .docx from markdown with SWS style canon."
    )
    ap.add_argument("output", help="Path to output .docx file")

    src = ap.add_mutually_exclusive_group(required=True)
    src.add_argument("--from-markdown", metavar="MD", help="Single markdown file")
    src.add_argument("--from-stitched", metavar="MD", help="Pre-stitched markdown file (alias)")
    src.add_argument("--from-drafts-dir", metavar="DIR", help="Directory of _drafts/*.md files")

    ap.add_argument("--profile", metavar="PROFILE",
                    help="Profile id (required with --from-drafts-dir)")
    args = ap.parse_args(argv)

    if args.from_drafts_dir and not args.profile:
        ap.error("--profile is required with --from-drafts-dir")

    # Load style definitions from references/docx-style.md.
    style_defs = _load_style_defs()

    # Read source markdown.
    if args.from_markdown:
        lines = _read_markdown_file(Path(args.from_markdown))
    elif args.from_stitched:
        lines = _read_markdown_file(Path(args.from_stitched))
    else:
        lines = _read_drafts_dir(Path(args.from_drafts_dir), args.profile)

    specs = _classify_lines(lines)
    _build_docx(specs, style_defs, Path(args.output))
    return 0


if __name__ == "__main__":
    sys.exit(main())
