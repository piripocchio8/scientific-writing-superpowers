"""Read a .docx file as plain text — SWS wrapper around python-docx.

Cycle #7 hotfix (D22): native Read tool fails on .docx, so SWS ships a
sanctioned reader. Drafting agents call this through
scripts/sws_python.sh "$PAPER_ROOT" sws_read_docx.py <args>.

CLI:
  sws_read_docx.py <file.docx>
  sws_read_docx.py <file.docx> --section "Introduction"
  sws_read_docx.py <file.docx> --paragraphs 3-7
  sws_read_docx.py <file.docx> --with-styles

Exit codes:
  0  ok
  2  file not found
  3  python-docx parse error (not a real docx, corrupt zip, etc.)
"""
from __future__ import annotations
import argparse
import sys
from pathlib import Path


def _is_heading(paragraph) -> bool:
    """Return True for paragraphs that act as headings.

    Primary signal: paragraph style name starts with 'Heading'.
    Fallback: the paragraph is a single-run, fully-bold line — covers
    docx files that hand-style headings without using the Heading styles.
    """
    style = getattr(paragraph.style, "name", "") or ""
    if style.startswith("Heading") or style == "Title":
        return True
    runs = list(paragraph.runs)
    if len(runs) == 1 and runs[0].bold and (paragraph.text or "").strip():
        return True
    return False


def _section_paragraphs(paragraphs, section_query: str):
    """Yield paragraphs that fall under the heading matching section_query.

    Match: case-insensitive substring on heading text. Capture starts at the
    matching heading (exclusive) and stops at the next heading of any level.
    """
    needle = section_query.strip().lower()
    in_section = False
    for p in paragraphs:
        if _is_heading(p):
            if in_section:
                # Reached the next heading — stop emitting.
                return
            if needle in (p.text or "").lower():
                in_section = True
                continue
        elif in_section:
            yield p


def _parse_paragraph_range(spec: str) -> tuple[int, int]:
    """Parse 'N-M' (1-indexed, inclusive). Single 'N' is treated as N-N."""
    if "-" in spec:
        lo_s, hi_s = spec.split("-", 1)
        lo, hi = int(lo_s), int(hi_s)
    else:
        lo = hi = int(spec)
    if lo < 1 or hi < lo:
        raise ValueError(f"invalid paragraph range: {spec!r}")
    return lo, hi


def _format_line(p, with_styles: bool) -> str:
    text = p.text or ""
    if with_styles:
        style = getattr(p.style, "name", "") or "Normal"
        return f"[{style}] {text}"
    return text


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        description="Read .docx as plain text (SWS wrapper around python-docx)."
    )
    ap.add_argument("file", help="Path to .docx file")
    ap.add_argument("--section", help="Print only paragraphs under matching heading")
    ap.add_argument("--paragraphs", help="1-indexed range, e.g. 5-12")
    ap.add_argument("--with-styles", action="store_true",
                    help="Prefix each line with [<style-name>]")
    args = ap.parse_args(argv)

    path = Path(args.file)
    if not path.is_file():
        print(f"sws_read_docx: file not found: {path}", file=sys.stderr)
        return 2

    try:
        from docx import Document
        doc = Document(str(path))
    except Exception as exc:  # python-docx raises a variety of errors on bad files
        print(f"sws_read_docx: parse error: {exc}", file=sys.stderr)
        return 3

    paragraphs = list(doc.paragraphs)

    if args.section:
        selected = list(_section_paragraphs(paragraphs, args.section))
    elif args.paragraphs:
        try:
            lo, hi = _parse_paragraph_range(args.paragraphs)
        except ValueError as exc:
            print(f"sws_read_docx: {exc}", file=sys.stderr)
            return 2
        selected = paragraphs[lo - 1:hi]
    else:
        selected = paragraphs

    for p in selected:
        print(_format_line(p, args.with_styles))
    return 0


if __name__ == "__main__":
    sys.exit(main())
